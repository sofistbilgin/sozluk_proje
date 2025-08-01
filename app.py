from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import joinedload
from docx import Document
from io import BytesIO
import re
from sqlalchemy import func

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///notdefteri.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# MODELLER
class MaddeBasi(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    baslik = db.Column(db.String(255), unique=True, nullable=False)
    anlamlar = db.relationship('Anlam', backref='maddebasi', cascade="all, delete-orphan", lazy=True)

class Anlam(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    maddebasi_id = db.Column(db.Integer, db.ForeignKey('madde_basi.id'), nullable=False)
    metin = db.Column(db.Text, nullable=False)
    kunyeler = db.Column(db.PickleType, default=[])

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/maddebasilar')
def maddebasilar():
    harf = request.args.get('harf', '').upper()
    query = MaddeBasi.query
    if harf:
        query = query.filter(MaddeBasi.baslik.ilike(f'{harf}%'))
    madde_listesi = query.order_by(MaddeBasi.baslik).all()
    return jsonify([{'id': m.id, 'baslik': m.baslik} for m in madde_listesi])

@app.route('/api/maddebasi/<int:id>')
def maddebasi(id):
    madde = MaddeBasi.query.options(joinedload(MaddeBasi.anlamlar)).get(id)
    if not madde:
        return jsonify({'error': 'Madde bulunamadı'}), 404
    return jsonify({
        'id': madde.id,
        'baslik': madde.baslik,
        'anlamlar': [
            {'id': a.id, 'metin': a.metin, 'kunyeler': a.kunyeler or []}
            for a in madde.anlamlar
        ]
    })

@app.route('/api/maddebasi', methods=['POST'])
def madde_ekle():
    veri = request.get_json()
    baslik = veri.get('baslik')
    if not baslik:
        return 'Başlık boş olamaz', 400
    if MaddeBasi.query.filter_by(baslik=baslik).first():
        return 'Bu başlık zaten var', 409
    madde = MaddeBasi(baslik=baslik)
    db.session.add(madde)
    db.session.commit()
    return jsonify({'id': madde.id, 'baslik': madde.baslik}), 201

@app.route('/api/madde-ekle-tam', methods=['POST'])
def madde_ve_anlam_ekle():
    veri = request.get_json()
    baslik = veri.get('baslik')
    anlamlar = veri.get('anlamlar', [])

    if not baslik or not anlamlar:
        return jsonify({'error': 'Başlık ve anlamlar zorunludur'}), 400

    if MaddeBasi.query.filter_by(baslik=baslik).first():
        return jsonify({'error': 'Bu başlık zaten var'}), 409

    yeni_madde = MaddeBasi(baslik=baslik)
    db.session.add(yeni_madde)
    db.session.flush()  # ID'yi almak için

    for anlam in anlamlar:
        metin = anlam.get('metin')
        kunyeler = anlam.get('kunyeler', [])
        yeni_anlam = Anlam(maddebasi_id=yeni_madde.id, metin=metin, kunyeler=kunyeler)
        db.session.add(yeni_anlam)

    db.session.commit()
    return jsonify({'message': 'Madde ve anlamlar eklendi', 'id': yeni_madde.id}), 201


@app.route('/api/maddebasi/<int:id>', methods=['PUT'])
def madde_guncelle(id):
    madde = MaddeBasi.query.get_or_404(id)
    veri = request.get_json()
    yeni_baslik = veri.get('baslik')
    if not yeni_baslik:
        return 'Yeni başlık boş olamaz', 400
    madde.baslik = yeni_baslik
    db.session.commit()
    return jsonify({'id': madde.id, 'baslik': madde.baslik})

@app.route('/api/maddebasi/<int:id>', methods=['DELETE'])
def madde_sil(id):
    madde = MaddeBasi.query.get_or_404(id)
    db.session.delete(madde)
    db.session.commit()
    return jsonify({'message': 'Silindi'})

@app.route('/api/anlam', methods=['POST'])
def anlam_ekle():
    veri = request.get_json()
    maddebasi_id = veri.get('maddebasi_id')
    metin = veri.get('metin')
    kunyeler = veri.get('kunyeler', [])
    if not maddebasi_id or not metin:
        return 'Geçersiz veri', 400
    anlam = Anlam(maddebasi_id=maddebasi_id, metin=metin, kunyeler=kunyeler)
    db.session.add(anlam)
    db.session.commit()
    return jsonify({'id': anlam.id})

@app.route('/api/anlam/<int:id>', methods=['PUT'])
def anlam_guncelle(id):
    anlam = Anlam.query.get_or_404(id)
    veri = request.get_json()
    anlam.metin = veri.get('metin', anlam.metin)
    anlam.kunyeler = veri.get('kunyeler', anlam.kunyeler)
    db.session.commit()
    return jsonify({'id': anlam.id})

@app.route('/api/anlam/<int:id>', methods=['DELETE'])
def anlam_sil(id):
    anlam = Anlam.query.get_or_404(id)
    db.session.delete(anlam)
    db.session.commit()
    return jsonify({'message': 'Silindi'})

# 📄 Word'e Aktar
@app.route('/api/maddebasi/<int:id>/export-word')
def maddebasi_export_word(id):
    madde = MaddeBasi.query.options(joinedload(MaddeBasi.anlamlar)).get(id)
    if not madde:
        return jsonify({'error': 'Madde bulunamadı'}), 404

    doc = Document()
    doc.add_heading(madde.baslik, level=1)

    for i, anlam in enumerate(madde.anlamlar, 1):
        doc.add_heading(f"Anlam {i}", level=2)
        doc.add_paragraph(anlam.metin)
        if anlam.kunyeler:
            doc.add_paragraph("Künyeler:", style='Intense Quote')
            for k in anlam.kunyeler:
                doc.add_paragraph(f"- {k}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{madde.baslik}_anlamlar.docx"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/api/export-all')
def export_all():
    madde_listesi = MaddeBasi.query.options(joinedload(MaddeBasi.anlamlar)).order_by(MaddeBasi.baslik).all()
    doc = Document()

    doc.add_heading('Tüm Madde Başlıkları ve Anlamları', 0)

    for madde in madde_listesi:
        doc.add_heading(madde.baslik, level=1)
        for anlam in madde.anlamlar:
            doc.add_paragraph(f"- {anlam.metin}", style='List Bullet')
            if anlam.kunyeler:
                for künye in anlam.kunyeler:
                    doc.add_paragraph(f"  📌 {künye}", style='List Bullet 2')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='tum_maddeler.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Kavram Haritası için yeni endpoint


@app.route('/api/maddebasi/<int:madde_id>/kavram-haritasi')
def kavram_haritasi(madde_id):
    secili = MaddeBasi.query.get_or_404(madde_id)
    tum_maddeler = MaddeBasi.query.all()

    madde_dict = {m.baslik.lower(): m for m in tum_maddeler if m.id != secili.id}

    nodes = []
    edges = []

    eklenen_node_idler = set()

    def anlamlari_al(madde):
        anlamlar = [a.metin for a in madde.anlamlar]
        return '\n'.join([f"{i+1}. {a}" for i, a in enumerate(anlamlar)]) if anlamlar else "Anlam bulunamadı"

    secili_anlamlar = anlamlari_al(secili)

    nodes.append({
        'id': f"madde-{secili.id}",
        'label': secili.baslik,
        'shape': 'box',
        'color': '#0d6efd',
        'font': {'color': 'white'},
        'title': secili_anlamlar
    })

    for anlam in secili.anlamlar:
        anlam_metin = anlam.metin.lower()
        for diger_baslik in madde_dict:
            if diger_baslik in anlam_metin:
                hedef = madde_dict[diger_baslik]
                hedef_id = f"madde-{hedef.id}"

                if hedef_id not in eklenen_node_idler:
                    nodes.append({
                        'id': hedef_id,
                        'label': hedef.baslik,
                        'color': '#198754',
                        'title': anlamlari_al(hedef)
                    })
                    eklenen_node_idler.add(hedef_id)

                edge = {'from': f"madde-{secili.id}", 'to': hedef_id}
                if edge not in edges:
                    edges.append(edge)

    return jsonify({'nodes': nodes, 'edges': edges})


@app.route('/api/maddebasi/by-baslik/<baslik>')
def maddebasi_by_baslik(baslik):
    madde = MaddeBasi.query.filter(func.lower(MaddeBasi.baslik) == baslik.lower()).first()
    if not madde:
        return jsonify({'error': 'Madde bulunamadı'}), 404

    return jsonify({
        'id': madde.id,
        'baslik': madde.baslik,
        'anlamlar': [
            {
                'id': anlam.id,
                'metin': anlam.metin,
                'kunyeler': anlam.kunyeler or []
            } for anlam in madde.anlamlar
        ]
    })



@app.route('/istatistikler')
def istatistikler():
    toplam_madde = db.session.query(Madde).count()
    toplam_anlam = db.session.query(Anlam).count()
    ortalama_anlam = round(toplam_anlam / toplam_madde, 2) if toplam_madde else 0
    return jsonify({
        'toplam_madde': toplam_madde,
        'toplam_anlam': toplam_anlam,
        'ortalama_anlam': ortalama_anlam
    })



if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run()
